"""발표용 데모 기획서 2종(DOCX) 생성 스크립트. 위반을 의도적으로 심어둔다."""
from pathlib import Path
from docx import Document

DEMO_DIR = Path(__file__).parent.parent / "demo"


def add_heading(doc: Document, text: str):
    doc.add_heading(text, level=2)


def build_mild():
    doc = Document()
    doc.add_heading("모여봐요 동네책방 — 서비스 기획서", level=1)

    add_heading(doc, "1. 서비스 개요")
    doc.add_paragraph(
        "동네 소규모 서점의 재고를 한눈에 검색하고 예약할 수 있는 지역 기반 도서 예약 플랫폼입니다. "
        "이용자는 앱을 통해 인근 서점의 도서 재고를 확인하고 방문 전 예약할 수 있습니다."
    )

    add_heading(doc, "2. 회원가입 및 개인정보 수집")
    doc.add_paragraph(
        "회원가입 시 이름, 이메일, 휴대폰번호를 수집하며, 회원가입 화면에서 개인정보 수집·이용에 대한 "
        "동의 체크박스를 별도로 마련하여 정보주체의 동의를 받습니다. 동의 화면에는 수집 항목, 목적, "
        "보유기간(회원 탈퇴 시까지)을 명시합니다."
    )
    doc.add_paragraph(
        "마케팅 정보 수신 동의는 필수 동의와 분리된 별도 체크박스로 구성하며 기본값은 선택하지 않은 "
        "상태(OFF)로 제공합니다."
    )

    add_heading(doc, "3. 위치 기반 서점 검색")
    doc.add_paragraph(
        "이용자가 '내 주변 서점 찾기' 버튼을 누르면 기기의 위치 권한 동의를 받은 후 대략적인 위치 정보를 "
        "1회성으로 조회하여 인근 서점 목록을 보여줍니다. 위치 정보는 서버에 저장하지 않고 즉시 폐기합니다."
    )

    add_heading(doc, "4. 도서 추천 기능")
    doc.add_paragraph(
        "이용자의 예약 이력을 기반으로 관심 있을 만한 도서를 추천합니다. 추천 로직은 마이페이지의 "
        "'추천 이유 보기' 메뉴에서 확인할 수 있으나, 현재 기획안에는 추천 결과에 대한 이의 제기나 "
        "추천 거부 절차에 대한 구체적인 안내 문구가 아직 포함되어 있지 않습니다."
    )

    add_heading(doc, "5. 데이터 보관 및 파기")
    doc.add_paragraph(
        "회원 탈퇴 시 개인정보는 즉시 파기하며, 관계 법령에 따라 보관이 필요한 거래 기록은 관련 법령이 "
        "정한 기간 동안 별도 보관 후 파기합니다."
    )

    doc.save(DEMO_DIR / "sample_mild.docx")


def build_severe():
    doc = Document()
    doc.add_heading("우리동네 중고마켓 — 서비스 기획서", level=1)

    add_heading(doc, "1. 서비스 개요")
    doc.add_paragraph(
        "동네 이웃 간 중고 물품 거래를 중개하는 하이퍼로컬 커머스 앱입니다. AI 기반 매칭과 "
        "위치 기반 추천으로 빠른 거래를 지원합니다."
    )

    add_heading(doc, "2. 간편 회원가입")
    doc.add_paragraph(
        "이름, 휴대폰번호, 이메일을 입력하면 별도의 동의 절차 없이 즉시 가입이 완료됩니다. "
        "회원가입 화면에는 개인정보 수집·이용에 대한 동의 문구나 체크박스가 없습니다."
    )
    doc.add_paragraph(
        "만 14세 미만 이용자도 생년월일 확인이나 법정대리인 동의 절차 없이 자유롭게 가입할 수 있습니다."
    )

    add_heading(doc, "3. 위치 기반 매칭")
    doc.add_paragraph(
        "가입 직후부터 이용자의 정확한 GPS 위치정보를 상시 수집하여 주변 판매자와 실시간 매칭합니다. "
        "수집된 위치정보는 제휴 광고 대행사에 제공되어 맞춤형 광고 타겟팅에 활용됩니다. "
        "이 정보 제공에 대해 이용자로부터 별도 동의를 받지 않습니다."
    )

    add_heading(doc, "4. AI 추천 및 신용점수 산정")
    doc.add_paragraph(
        "이용자의 거래 이력, 채팅 응답 속도, 매너 평가를 AI가 자동 분석하여 '신뢰도 점수'를 산정하고, "
        "이 점수가 낮으면 별도 고지 없이 검색 노출 순위를 자동으로 낮춥니다. 이용자는 이 자동화된 "
        "결정에 대해 설명을 요구하거나 이의를 제기할 수 있는 절차를 안내받지 못합니다."
    )
    doc.add_paragraph(
        "또한 이용자들의 채팅 대화 내용을 신뢰도 점수 산정 AI 모델의 추가 학습 데이터로 활용할 "
        "예정이며, 이에 대한 별도 동의 절차는 기획에 포함되어 있지 않습니다."
    )

    add_heading(doc, "5. 데이터 보관")
    doc.add_paragraph(
        "수집된 개인정보 및 위치정보의 보유기간에 대한 안내는 이용약관 및 개인정보 처리방침에 "
        "별도로 명시하지 않습니다."
    )

    doc.save(DEMO_DIR / "sample_severe.docx")


if __name__ == "__main__":
    DEMO_DIR.mkdir(parents=True, exist_ok=True)
    build_mild()
    build_severe()
    print("데모 문서 생성 완료:", list(DEMO_DIR.glob("*.docx")))
