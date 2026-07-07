"""스크리닝 파이프라인 오케스트레이터.

업로드 → Pass 1 (OCR/텍스트 추출 + 리스크 포인트·인벤토리) → Clova 쿼리 임베딩 →
FAISS 검색 → Pass 2 (RAG 판정) → ScreeningReport

Pass 1은 항상 Gemini(멀티모달)가 전담한다. Pass 2는 pass2_client로 주입받아
HyperCLOVA X / Gemini 중 무엇이든 사용할 수 있다 (modules.llm_provider 참조).
"""
import logging
from typing import Callable, Optional

import fitz  # pymupdf
import numpy as np

from modules.schemas import (
    Pass1Output, ScreeningReport, RiskPoint, RetrievedChunk, CrossExamVerdict
)
from modules.gemini_client import GeminiClient
from modules.embedder_clova import ClovaEmbedder
from modules.vectorstore import VectorStore

logger = logging.getLogger(__name__)

# 총 주입 예산: Pass 2에 전달할 최대 청크 수 (Parent 확장 후)
MAX_CONTEXT_CHUNKS = 12
# 포인트당 보장 근거 수 (전역 컷에서 밀려나도 최소한의 근거는 확보)
MIN_CHUNKS_PER_POINT = 3
# 텍스트 레이어가 이 길이 이상이면 "스캔본 아님"으로 판단해 OCR을 생략
TEXT_LAYER_MIN_CHARS = 30


def pdf_to_images(file_bytes: bytes, dpi: int = 200) -> list[bytes]:
    """PDF를 페이지별 PNG 이미지로 변환.

    Args:
        file_bytes: PDF 파일 바이트
        dpi: 렌더링 DPI (기본 200)

    Returns:
        페이지별 PNG 바이트 목록
    """
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    images = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        # DPI 기반 확대 행렬
        zoom = dpi / 72
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat)
        images.append(pix.tobytes("png"))

    doc.close()
    return images


def extract_pdf_text_layer(file_bytes: bytes) -> Optional[str]:
    """PDF에 추출 가능한 텍스트 레이어가 있으면 페이지 구분자와 함께 반환.

    스캔본(이미지만 있는 PDF)이면 None을 반환해 OCR 경로로 넘긴다.
    """
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    try:
        pages_text = [doc[i].get_text().strip() for i in range(len(doc))]
    finally:
        doc.close()

    total_chars = sum(len(t) for t in pages_text)
    if total_chars < TEXT_LAYER_MIN_CHARS:
        return None

    return "\n\n".join(f"=== PAGE {i + 1} ===\n{t}" for i, t in enumerate(pages_text))


def extract_docx_text(file_bytes: bytes) -> str:
    """DOCX에서 본문·표 텍스트를 순서대로 추출."""
    import io
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _select_context_chunks(
    point_chunks: dict[str, list[RetrievedChunk]],
) -> list[RetrievedChunk]:
    """포인트당 top-N 근거를 우선 보장한 뒤, 남는 예산을 전역 점수순으로 채운다.

    v1은 전역 풀을 점수순으로 컷했기 때문에 특정 포인트의 근거가 전부
    밀려나 판정 보류로 빠지는 문제가 있었다. 포인트별 보장 슬롯을 먼저
    채워 이를 방지한다.
    """
    seen_parents: set[str] = set()
    seen_chunk_ids: set[str] = set()
    reserved: list[RetrievedChunk] = []

    def _dedup_add(chunk: RetrievedChunk) -> bool:
        parent_key = chunk.meta.parent_id or chunk.chunk_id
        if parent_key in seen_parents or chunk.chunk_id in seen_chunk_ids:
            return False
        seen_parents.add(parent_key)
        seen_chunk_ids.add(chunk.chunk_id)
        return True

    # 1. 포인트별 보장 슬롯
    for chunks in point_chunks.values():
        ranked = sorted(chunks, key=lambda c: c.score, reverse=True)
        taken = 0
        for chunk in ranked:
            if taken >= MIN_CHUNKS_PER_POINT:
                break
            if _dedup_add(chunk):
                reserved.append(chunk)
                taken += 1

    # 2. 남는 예산은 전역 점수순으로 채움
    remaining_budget = MAX_CONTEXT_CHUNKS - len(reserved)
    if remaining_budget > 0:
        overflow_pool = [
            c for chunks in point_chunks.values() for c in chunks
        ]
        overflow_pool.sort(key=lambda c: c.score, reverse=True)
        for chunk in overflow_pool:
            if remaining_budget <= 0:
                break
            if _dedup_add(chunk):
                reserved.append(chunk)
                remaining_budget -= 1

    reserved.sort(key=lambda c: c.score, reverse=True)
    return reserved[:MAX_CONTEXT_CHUNKS]


def run_screening(
    file_bytes: bytes,
    mime_type: str,
    pass1_client: GeminiClient,
    pass2_client,
    embedder: ClovaEmbedder,
    vectorstore: VectorStore,
    law_ids: Optional[list[str]] = None,
    progress_callback: Optional[Callable[[str, str], None]] = None,
) -> tuple[ScreeningReport, Pass1Output]:
    """전체 스크리닝 파이프라인 실행.

    Args:
        file_bytes: 업로드된 파일 바이트
        mime_type: 파일 MIME 타입 (application/pdf, image/png, image/jpeg,
            application/vnd.openxmlformats-officedocument.wordprocessingml.document)
        pass1_client: Pass 1(OCR·추출) 클라이언트 — 항상 Gemini
        pass2_client: Pass 2(RAG 판정) 클라이언트 — Gemini 또는 HCX
            (.run_pass2(full_text, risk_points, chunks) -> ScreeningReport 인터페이스만 요구)
        embedder: Clova 임베딩 클라이언트
        vectorstore: FAISS 벡터 저장소
        law_ids: 검토 대상 법령 도메인 필터 (None/빈 리스트면 전체)
        progress_callback: 진행 상황 콜백 (stage_name, detail)

    Returns:
        (ScreeningReport, Pass1Output) — Pass1Output은 인벤토리·재검사 루프에 재사용
    """
    def _progress(stage: str, detail: str = ""):
        if progress_callback:
            progress_callback(stage, detail)
        logger.info(f"[{stage}] {detail}")

    # ── 1단계: 문서 준비 (텍스트 직접 추출 우선, 안 되면 이미지 OCR) ──
    _progress("준비", "문서 형식을 확인하는 중...")

    docx_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    use_text_path = False
    prepared_text = ""
    images: list[bytes] = []

    if mime_type == docx_mime:
        prepared_text = extract_docx_text(file_bytes)
        use_text_path = True
        _progress("준비", f"DOCX 텍스트 추출 완료 ({len(prepared_text)}자)")
    elif mime_type == "application/pdf":
        text_layer = extract_pdf_text_layer(file_bytes)
        if text_layer is not None:
            prepared_text = text_layer
            use_text_path = True
            _progress("준비", f"PDF 텍스트 레이어 직접 추출 완료 ({len(prepared_text)}자)")
        else:
            images = pdf_to_images(file_bytes)
            _progress("준비", f"스캔본으로 판단 — {len(images)}페이지 이미지 변환 완료")
    elif mime_type in ("image/png", "image/jpeg", "image/jpg"):
        images = [file_bytes]
        _progress("준비", "이미지 1장 로드 완료")
    else:
        raise ValueError(f"지원하지 않는 파일 형식: {mime_type}")

    # ── 2단계: Pass 1 (텍스트 추출/OCR + 리스크 포인트·인벤토리) ──
    _progress("리스크 추출", "Pass 1: 리스크 포인트 및 개인정보 인벤토리 식별 중...")
    if use_text_path:
        pass1_output: Pass1Output = pass1_client.run_pass1_text(prepared_text)
    else:
        pass1_output = pass1_client.run_pass1(images)

    risk_count = len(pass1_output.risk_points)
    _progress(
        "리스크 추출",
        f"완료 — {risk_count}개 리스크 포인트, {len(pass1_output.inventory)}개 인벤토리 항목 식별",
    )

    if risk_count == 0:
        _progress("완료", "리스크 포인트가 없습니다.")
        return ScreeningReport(
            findings=[],
            inconclusive_points=[],
            document_summary=pass1_output.full_text[:500] + "...",
            disclaimer="본 결과는 AI 사전 스크리닝이며 법률 자문을 대체하지 않습니다.",
        ), pass1_output

    # ── 3단계: 쿼리 임베딩 + FAISS 검색 (포인트별 근거 보관) ──
    _progress("법령 검색", "관련 법령 조문 검색 중...")

    point_chunks: dict[str, list[RetrievedChunk]] = {}
    no_result_points: list[str] = []  # 근거 미검색 포인트
    law_id_filter = set(law_ids) if law_ids else None

    for rp in pass1_output.risk_points:
        try:
            qvec = embedder.embed_query(rp.query_text)
            results = vectorstore.search(qvec, k=8, threshold=0.35)

            if law_id_filter is not None:
                results = [r for r in results if r.meta.law_id in law_id_filter]

            if not results:
                no_result_points.append(rp.point_id)
                _progress("법령 검색", f"{rp.point_id}: 근거 미검색")
            else:
                expanded = [vectorstore.expand_to_parent(c) for c in results]
                point_chunks[rp.point_id] = expanded
                _progress("법령 검색", f"{rp.point_id}: {len(results)}건 검색")
        except Exception as e:
            logger.warning(f"{rp.point_id} 임베딩/검색 실패: {e}")
            no_result_points.append(rp.point_id)

    context_chunks = _select_context_chunks(point_chunks)
    _progress("법령 검색", f"총 {len(context_chunks)}개 근거 청크 선별 완료 (포인트별 보장 적용)")

    # ── 4단계: Pass 2 (RAG 판정) ──
    _progress("판정", "Pass 2: 위반 판정 및 개선안 생성 중...")

    # 근거 미검색 포인트는 Pass 2에서 제외
    active_risk_points = [
        rp for rp in pass1_output.risk_points
        if rp.point_id not in no_result_points
    ]

    if not active_risk_points and not context_chunks:
        _progress("완료", "검색된 근거가 없어 판정을 수행할 수 없습니다.")
        return ScreeningReport(
            findings=[],
            inconclusive_points=[rp.point_id for rp in pass1_output.risk_points],
            document_summary=pass1_output.full_text[:500] + "...",
            disclaimer="본 결과는 AI 사전 스크리닝이며 법률 자문을 대체하지 않습니다.",
        ), pass1_output

    logger.debug(
        "LLM 컨텍스트: %s",
        "\n\n".join(f"[chunk_id: {c.chunk_id}]\n{c.meta.text}" for c in context_chunks),
    )

    report = pass2_client.run_pass2(
        full_text=pass1_output.full_text,
        risk_points=active_risk_points,
        chunks=context_chunks,
    )

    # ── 5단계: 교차검증 (반론 에이전트) — 오탐 방어선 ──
    if report.findings:
        _progress("교차검증", f"{len(report.findings)}건 판정에 대한 반론 검토 중...")
        try:
            cross_results = pass2_client.run_cross_exam(
                full_text=pass1_output.full_text,
                findings=report.findings,
                chunks=context_chunks,
            )
        except Exception as e:
            logger.warning(f"교차검증 실패, 원 판정을 그대로 유지합니다: {e}")
            cross_results = []

        cross_map = {r.point_id: r for r in cross_results}
        surviving_findings = []
        for finding in report.findings:
            cx = cross_map.get(finding.point_id)
            if cx is None:
                surviving_findings.append(finding)
                continue

            finding.cross_exam = cx
            if cx.verdict == CrossExamVerdict.OVERTURN:
                _progress("교차검증", f"{finding.point_id} 반론 에이전트: \"{cx.rebuttal}\" → 인용·판정 철회")
                if finding.point_id not in report.inconclusive_points:
                    report.inconclusive_points.append(finding.point_id)
            elif cx.verdict == CrossExamVerdict.WEAKEN:
                finding.confidence = min(finding.confidence, 0.4)
                _progress("교차검증", f"{finding.point_id} 반론 에이전트: \"{cx.rebuttal}\" → 일부 인정·확신도 하향")
                surviving_findings.append(finding)
            else:
                _progress("교차검증", f"{finding.point_id} 반론 에이전트: \"{cx.rebuttal}\" → 기각·판정 유지")
                surviving_findings.append(finding)

        report = ScreeningReport(
            findings=surviving_findings,
            inconclusive_points=report.inconclusive_points,
            document_summary=report.document_summary,
            disclaimer=report.disclaimer,
        )

    # 근거 미검색 포인트를 inconclusive에 병합
    existing_inconclusive = set(report.inconclusive_points)
    for pid in no_result_points:
        if pid not in existing_inconclusive:
            report.inconclusive_points.append(pid)

    _progress("완료", f"분석 완료 — {len(report.findings)}건 위반, {len(report.inconclusive_points)}건 보류")
    return report, pass1_output
