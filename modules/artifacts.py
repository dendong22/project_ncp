"""컴플라이언스 산출물(개인정보 처리방침 초안·동의 문구) 생성.

스크리닝의 부산물인 개인정보 인벤토리를 재활용해 문서 초안을 생성한다.
"""
import io
from pathlib import Path

from docx import Document

from modules.schemas import InventoryItem

PROMPTS_DIR = Path(__file__).parent.parent / "prompts"


def _inventory_to_text(inventory: list[InventoryItem]) -> str:
    if not inventory:
        return "(추출된 개인정보 처리 항목 없음)"
    lines = []
    for it in inventory:
        lines.append(
            f"- {it.item} | 분류: {it.category} | 목적: {it.purpose} | "
            f"보유기간: {it.retention} | 제3자 제공: {it.third_party or '없음'} | "
            f"수명주기: {', '.join(it.lifecycle_stages) or '-'}"
        )
    return "\n".join(lines)


def generate_privacy_policy(pass2_client, inventory: list[InventoryItem]) -> str:
    """개인정보 처리방침 초안(마크다운) 생성."""
    prompt = (PROMPTS_DIR / "policy_draft.txt").read_text(encoding="utf-8")
    user_message = f"<inventory>\n{_inventory_to_text(inventory)}\n</inventory>"
    return pass2_client.generate_text(prompt, user_message)


def generate_consent_copy(pass2_client, inventory: list[InventoryItem]) -> str:
    """동의 UI 문구 초안(마크다운) 생성."""
    prompt = (PROMPTS_DIR / "consent_copy.txt").read_text(encoding="utf-8")
    user_message = f"<inventory>\n{_inventory_to_text(inventory)}\n</inventory>"
    return pass2_client.generate_text(prompt, user_message)


def markdown_to_docx_bytes(title: str, text: str) -> bytes:
    """단순 마크다운(헤더·리스트)을 DOCX 바이트로 변환."""
    doc = Document()
    doc.add_heading(title, level=1)
    for raw_line in text.split("\n"):
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith(("- ", "* ")):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
